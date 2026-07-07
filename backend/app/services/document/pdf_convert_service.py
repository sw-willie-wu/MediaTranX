"""
PDF / document conversion service.
Supports: PDF -> TXT / Markdown / Images(zip)
           DOCX -> TXT / Markdown
"""
import io
import logging
import zipfile
from pathlib import Path
from typing import Callable, Optional
from uuid import uuid4

from app.services.files.file_service import FileService
from app.workers.task_manager import TaskManager

logger = logging.getLogger(__name__)

TASK_TYPE_DOCUMENT_PDF_CONVERT = "document.pdf_convert"


class DocumentPdfConvertService:
    """Document conversion service for PDF/DOCX to TXT, Markdown, or images."""

    def __init__(self, file_service: FileService, task_manager: TaskManager):
        self._file_service = file_service
        self._task_manager = task_manager
        self._task_manager.register_handler(
            TASK_TYPE_DOCUMENT_PDF_CONVERT, self._handle_task,
            output_policy="results",
        )
        logger.info("DocumentPdfConvertService initialized")

    async def submit(
        self,
        file_id: str,
        output_format: str = "txt",
        suppress_results: Optional[bool] = None,
    ) -> str:
        file_info = self._file_service.require_file(file_id)
        params = {
            "file_id": file_id,
            "output_format": output_format,
        }
        task_id = await self._task_manager.submit(
            TASK_TYPE_DOCUMENT_PDF_CONVERT, params, suppress_results=suppress_results
        )
        logger.info(f"Document PDF convert task submitted: {task_id}")
        return task_id

    def _handle_task(self, params: dict, progress_callback: Callable[[float, str], None]) -> dict:
        return self._execute(params, progress_callback)

    def _execute(self, params: dict, progress_callback: Callable[[float, str], None]) -> dict:
        file_id = params["file_id"]
        file_info = self._file_service.require_file(file_id)

        src_path = file_info.file_path
        src_ext = Path(file_info.original_filename).suffix.lower()
        output_format = params.get("output_format", "txt")
        output_file_id = str(uuid4())
        stem = Path(file_info.original_filename).stem

        ext_map = {"txt": "txt", "md": "md", "images": "zip"}
        ext = ext_map.get(output_format, "txt")

        final_filename = f"{stem}_converted.{ext}"

        output_dir = self._file_service.output_dir
        output_path = output_dir / final_filename

        progress_callback(0.05, "task.progress.reading_document")

        if output_format == "images":
            if src_ext != ".pdf":
                raise ValueError("Only PDF format supports conversion to images")
            self._pdf_to_images(src_path, output_path, progress_callback)
        elif src_ext == ".pdf":
            text = self._extract_pdf_text(src_path, progress_callback)
            content = self._maybe_to_md(text) if output_format == "md" else text
            output_path.write_text(content, encoding="utf-8")
        elif src_ext in (".docx", ".doc"):
            text = self._extract_docx_text(src_path, progress_callback)
            content = self._maybe_to_md(text) if output_format == "md" else text
            output_path.write_text(content, encoding="utf-8")
        else:
            # Plain text files: copy directly
            content = src_path.read_text(encoding="utf-8", errors="replace")
            output_path.write_text(content, encoding="utf-8")

        output_info = self._file_service.register_output(
            file_id=output_file_id,
            file_path=output_path,
            original_filename=final_filename,
        )

        progress_callback(1.0, "task.progress.convert_complete")
        return {
            "output_file_id": output_file_id,
            "output_filename": output_info.filename,
        }

    def _extract_pdf_text(self, path: Path, progress_callback) -> str:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        parts = []
        total = len(reader.pages)
        for i, page in enumerate(reader.pages):
            parts.append(page.extract_text() or "")
            progress_callback(0.1 + (i + 1) / total * 0.8, f"task.progress.extracting_page|{i+1}|{total}")
        return "\n\n".join(p.strip() for p in parts if p.strip())

    def _extract_docx_text(self, path: Path, progress_callback) -> str:
        import docx
        doc = docx.Document(str(path))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        progress_callback(0.9, "task.progress.text_extraction_complete")
        return "\n".join(lines)

    def _maybe_to_md(self, text: str) -> str:
        return text  # Plain text is readable enough; return as-is

    def _pdf_to_images(self, src_path: Path, output_path: Path, progress_callback):
        import pypdfium2
        doc = pypdfium2.PdfDocument(str(src_path))
        total = len(doc)
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, page in enumerate(doc):
                bitmap = page.render(scale=2.0)
                img = bitmap.to_pil()
                img_buf = io.BytesIO()
                img.save(img_buf, format="PNG")
                zf.writestr(f"page_{i + 1:03d}.png", img_buf.getvalue())
                progress_callback(0.1 + (i + 1) / total * 0.85, f"task.progress.rendering_page|{i+1}|{total}")
        doc.close()
        output_path.write_bytes(buf.getvalue())
